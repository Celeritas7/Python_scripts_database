"""
IPYNB to Formatted Word Converter (Windows)
============================================
Converts notebooks to Word with formatting matching VBA script.

Pipeline: ipynb → Markdown (via nbconvert) → DOCX (via Pandoc/pypandoc) → Format

Formatting (matches FormatFullDocument_GridTable4 VBA):
  1. Narrow margins (0.5 inch)
  2. Tables: Grid Table 4 style, Cambria 11pt, centered, autofit
  3. Images: Lock aspect ratio, scale 70%, centered

Usage:
  Double-click to run. Processes all .ipynb in script's folder.

Output:
  Word_Outputs/ folder with formatted .docx files

Requirements:
  - pip install nbconvert pypandoc python-docx nbformat
"""

import sys
import tempfile
from pathlib import Path

# Check dependencies
missing = []
try:
    import nbconvert
    from nbconvert import MarkdownExporter
except ImportError:
    missing.append("nbconvert")

try:
    import pypandoc
except ImportError:
    missing.append("pypandoc")

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    missing.append("python-docx")

try:
    import nbformat
except ImportError:
    missing.append("nbformat")

if missing:
    print("=" * 50)
    print("ERROR: Missing packages!")
    print(f"Run: pip install {' '.join(missing)}")
    print("=" * 50)
    input("Press Enter to exit...")
    sys.exit(1)


# =============================================================================
# CONFIGURATION
# =============================================================================

IMAGE_SCALE = 0.70  # 70%

# Heading styles
HEADING_STYLES = {
    'Title': {'size': 26, 'bold': False, 'color': (0x17, 0x36, 0x5D)},
    'Heading 1': {'size': 14, 'bold': True, 'color': (0x36, 0x5F, 0x91)},
    'Heading 2': {'size': 13, 'bold': True, 'color': (0x4F, 0x81, 0xBD)},
    'Heading 3': {'size': 11, 'bold': True, 'color': (0x4F, 0x81, 0xBD)},
}


# =============================================================================
# FORMATTING FUNCTIONS
# =============================================================================

def set_narrow_margins(doc):
    """Set margins to Narrow (0.5 inches) - VBA Step 1."""
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)


def apply_heading_styles(doc):
    """Apply custom heading styles."""
    for style_name, fmt in HEADING_STYLES.items():
        try:
            style = doc.styles[style_name]
            font = style.font
            font.size = Pt(fmt['size'])
            font.bold = fmt['bold']
            font.color.rgb = RGBColor(*fmt['color'])
        except KeyError:
            pass


def remove_trailing_paragraph_marks(doc):
    """Remove trailing ¶ symbols from headings."""
    for para in doc.paragraphs:
        if para.style and para.style.name in ['Title', 'Heading 1', 'Heading 2', 'Heading 3']:
            for run in para.runs:
                if run.text:
                    run.text = run.text.rstrip('\n\r \t')


def set_cell_shading(cell, hex_color):
    """Set background shading for a cell."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), hex_color)
    shading.set(qn('w:val'), 'clear')
    tcPr = cell._tc.get_or_add_tcPr()
    # Remove existing shading
    existing = tcPr.find(qn('w:shd'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(shading)


def format_tables(doc):
    """
    Format tables to match Grid Table 4 style - VBA Step 2.
    - Cambria font, 11pt
    - Header row: black background, white bold text
    - Banded rows (alternating gray)
    - AutoFit, centered
    """
    formatted = 0
    
    for table in doc.tables:
        try:
            # Try to apply built-in style first
            try:
                table.style = 'Grid Table 4'
            except:
                pass  # Style may not exist, we'll format manually
            
            # AutoFit to window - set table width to 100%
            table.autofit = True
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Set table to autofit window width
            tbl = table._tbl
            tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
            
            # Set table width to 100% (5000 = 100% in fifths of a percent)
            tblW = tblPr.find(qn('w:tblW'))
            if tblW is None:
                tblW = OxmlElement('w:tblW')
                tblPr.append(tblW)
            tblW.set(qn('w:w'), '5000')
            tblW.set(qn('w:type'), 'pct')
            
            if tbl.tblPr is None:
                tbl.insert(0, tblPr)
            
            # Process each row
            for row_idx, row in enumerate(table.rows):
                # Center row
                row_tr = row._tr
                trPr = row_tr.get_or_add_trPr()
                jc = trPr.find(qn('w:jc'))
                if jc is None:
                    jc = OxmlElement('w:jc')
                    trPr.append(jc)
                jc.set(qn('w:val'), 'center')
                
                for cell in row.cells:
                    # Set font to Cambria 11pt for all cells
                    for para in cell.paragraphs:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in para.runs:
                            run.font.name = 'Cambria'
                            run.font.size = Pt(11)
                    
                    # Vertical alignment - center
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    
                    # Header row (row 0): black background, white bold text
                    if row_idx == 0:
                        set_cell_shading(cell, '000000')
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                                run.font.bold = True
                    # Banded rows: alternating gray (even data rows)
                    elif row_idx % 2 == 0:
                        set_cell_shading(cell, 'F2F2F2')
            
            # Add borders
            tblBorders = OxmlElement('w:tblBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:color'), '000000')
                tblBorders.append(border)
            
            existing_borders = tblPr.find(qn('w:tblBorders'))
            if existing_borders is not None:
                tblPr.remove(existing_borders)
            tblPr.append(tblBorders)
            
            formatted += 1
            
        except Exception as e:
            print(f"      [WARN] Table error: {e}")
    
    return formatted


def process_images(doc):
    """
    Process all images - VBA Steps 3 & 4.
    - Lock aspect ratio
    - Scale to 70%
    - Center paragraph
    """
    processed = 0
    
    for shape in doc.inline_shapes:
        try:
            # Access the underlying XML
            inline = shape._inline
            
            # Find extent element (contains width/height in EMUs)
            extent = inline.find(qn('wp:extent'))
            if extent is not None:
                cx = int(extent.get('cx'))
                cy = int(extent.get('cy'))
                
                # Scale to 70%
                new_cx = int(cx * IMAGE_SCALE)
                new_cy = int(cy * IMAGE_SCALE)
                
                extent.set('cx', str(new_cx))
                extent.set('cy', str(new_cy))
                
                # Also update the graphic extent (a:ext elements)
                for ext in inline.iter(qn('a:ext')):
                    if ext.get('cx') and ext.get('cy'):
                        ext.set('cx', str(new_cx))
                        ext.set('cy', str(new_cy))
            
            # Center the paragraph containing this image
            # Navigate up to find the paragraph element
            parent = inline.getparent()
            while parent is not None:
                if parent.tag == qn('w:p'):
                    # Found the paragraph, set alignment
                    pPr = parent.find(qn('w:pPr'))
                    if pPr is None:
                        pPr = OxmlElement('w:pPr')
                        parent.insert(0, pPr)
                    
                    jc = pPr.find(qn('w:jc'))
                    if jc is None:
                        jc = OxmlElement('w:jc')
                        pPr.append(jc)
                    jc.set(qn('w:val'), 'center')
                    break
                parent = parent.getparent()
            
            processed += 1
            
        except Exception as e:
            pass
    
    return processed


def format_docx(docx_path: Path) -> dict:
    """Apply all formatting to the docx file."""
    stats = {'images': 0, 'tables': 0}
    
    try:
        doc = Document(str(docx_path))
        
        # 1. Set narrow margins (0.5 inch)
        set_narrow_margins(doc)
        
        # 2. Apply heading styles
        apply_heading_styles(doc)
        
        # 3. Remove trailing ¶ from headings
        remove_trailing_paragraph_marks(doc)
        
        # 4. Format tables (Grid Table 4 style)
        stats['tables'] = format_tables(doc)
        
        # 5. Process images (scale 70%, center)
        stats['images'] = process_images(doc)
        
        # Save
        doc.save(str(docx_path))
        
        return stats
        
    except Exception as e:
        print(f"      [ERROR] Formatting failed: {e}")
        import traceback
        traceback.print_exc()
        return stats


def convert_notebook(ipynb_path: Path, output_folder: Path) -> bool:
    """Convert ipynb → Markdown → DOCX with formatting."""
    print(f"\n  Processing: {ipynb_path.name}")
    
    docx_path = output_folder / f"{ipynb_path.stem}.docx"
    
    try:
        # Step 1: Read notebook
        print("    [1/3] Reading notebook...")
        with open(ipynb_path, 'r', encoding='utf-8') as f:
            notebook = nbformat.read(f, as_version=4)
        
        # Step 2: Convert to Markdown then DOCX (preserves syntax highlighting via Pandoc)
        print("    [2/3] Converting to Word (with syntax highlighting)...")
        md_exporter = MarkdownExporter()
        md_exporter.template_name = 'classic'
        (md_content, resources) = md_exporter.from_notebook_node(notebook)
        
        # Write markdown + any extracted outputs (images, etc.) into a temporary working directory
        tmp_dir = Path(tempfile.mkdtemp(prefix="ipynb2docx_"))
        tmp_md_path = tmp_dir / "notebook.md"
        tmp_md_path.write_text(md_content, encoding="utf-8")
        
        # nbconvert keeps binary outputs in resources['outputs'] as {filename: bytes}
        for fname, data in (resources.get("outputs") or {}).items():
            out_path = tmp_dir / fname
            out_path.parent.mkdir(parents=True, exist_ok=True)
            # data may be bytes or str; handle both
            if isinstance(data, bytes):
                out_path.write_bytes(data)
            else:
                out_path.write_text(str(data), encoding="utf-8")
        
        # nbconvert may also store attachments in resources['attachments'] as {cell_id: {filename: base64str}}
        # We don't try to rewrite markdown links here; but we still materialize attachments in case they are referenced.
        import base64
        for _cell_id, att_map in (resources.get("attachments") or {}).items():
            for fname, b64data in (att_map or {}).items():
                out_path = tmp_dir / fname
                out_path.parent.mkdir(parents=True, exist_ok=True)
                try:
                    out_path.write_bytes(base64.b64decode(b64data))
                except Exception:
                    # Fallback: write raw text if not valid base64
                    out_path.write_text(str(b64data), encoding="utf-8")
        
        try:
            pypandoc.convert_file(
                str(tmp_md_path),
                'docx',
                outputfile=str(docx_path),
                extra_args=[
                    '--from=markdown',
                    '--standalone',
                    '--highlight-style=tango',
                    f'--resource-path={str(tmp_dir)}',
                ],
            )
        finally:
            # Cleanup temp directory
            try:
                for p in sorted(tmp_dir.rglob("*"), reverse=True):
                    if p.is_file():
                        p.unlink(missing_ok=True)
                    else:
                        p.rmdir()
                tmp_dir.rmdir()
            except Exception:
                pass
        # Step 3: Apply formatting
        print("    [3/3] Formatting (margins, tables, images)...")
        stats = format_docx(docx_path)
        print(f"      Done: {stats['tables']} tables, {stats['images']} images")
        
        print(f"    [SUCCESS] → {docx_path.name}")
        return True
        
    except Exception as e:
        print(f"    [ERROR] Conversion failed: {e}")
        import traceback
        traceback.print_exc()
        return False


def main():
    print("=" * 60)
    print("    IPYNB → Formatted Word Converter")
    print("    (Grid Table 4 + 70% Images + Narrow Margins)")
    print("=" * 60)
    
    script_folder = Path(__file__).parent.resolve()
    print(f"\nScanning: {script_folder}")
    
    notebooks = [f for f in script_folder.glob("*.ipynb") 
                 if ".ipynb_checkpoints" not in str(f)]
    
    if not notebooks:
        print("\n[INFO] No .ipynb files found in this folder.")
        input("\nPress Enter to exit...")
        return
    
    print(f"Found {len(notebooks)} notebook(s)")
    
    output_folder = script_folder / "Word_Outputs"
    output_folder.mkdir(exist_ok=True)
    print(f"Output: {output_folder}")
    print("-" * 60)
    
    success = 0
    failed = 0
    
    for nb in notebooks:
        if convert_notebook(nb, output_folder):
            success += 1
        else:
            failed += 1
    
    print("\n" + "=" * 60)
    print(f"COMPLETE! Success: {success} | Failed: {failed}")
    print("=" * 60)
    
    input("\nPress Enter to exit...")


if __name__ == "__main__":
    main()
