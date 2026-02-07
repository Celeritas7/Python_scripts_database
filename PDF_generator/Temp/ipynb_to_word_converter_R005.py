"""
IPYNB to Formatted Word Converter (Windows)
============================================
Converts notebooks to Word with proper formatting.

Pipeline: ipynb → HTML (via nbconvert) → DOCX (via pypandoc) → Format

Formatting applied:
  - Narrow margins (1.27 cm / 0.5 inch)
  - Images: centered, locked aspect ratio, scaled to 70%
  - Removes trailing ¶ symbols from headings
  - Skips gray placeholder images (matplotlib toolbar icons)
  - Tables: Grid Table 4 style (black header, alternating rows)
  - Custom Title/Heading styles

Usage:
  Double-click to run. Processes all .ipynb in script's folder.

Output:
  Word_Outputs/ folder with formatted .docx files

Requirements:
  - pip install nbconvert pypandoc python-docx nbformat lxml
"""

import sys
import tempfile
from pathlib import Path

# Check dependencies
missing = []
try:
    import nbconvert
    from nbconvert import HTMLExporter
except ImportError:
    missing.append("nbconvert")

try:
    import pypandoc
except ImportError:
    missing.append("pypandoc")

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsmap
    from docx.oxml import OxmlElement
except ImportError:
    missing.append("python-docx")

try:
    import nbformat
except ImportError:
    missing.append("nbformat")

try:
    from lxml import etree
except ImportError:
    missing.append("lxml")

if missing:
    print("=" * 50)
    print("ERROR: Missing packages!")
    print(f"Run: pip install {' '.join(missing)}")
    print("=" * 50)
    input("Press Enter to exit...")
    sys.exit(1)


# =============================================================================
# CONFIGURATION
# =============================================================================

# Margins (Narrow = 1.27 cm = 0.5 inch)
MARGIN_CM = 1.27

# Image scaling (70%)
IMAGE_SCALE = 0.70

# Heading styles (from your sample document)
HEADING_STYLES = {
    'Title': {
        'size': 26,
        'bold': False,
        'color': (0x17, 0x36, 0x5D),  # Dark navy #17365D
    },
    'Heading 1': {
        'size': 14,
        'bold': True,
        'color': (0x36, 0x5F, 0x91),  # Medium blue #365F91
    },
    'Heading 2': {
        'size': 13,
        'bold': True,
        'color': (0x4F, 0x81, 0xBD),  # Light blue #4F81BD
    },
    'Heading 3': {
        'size': 11,
        'bold': True,
        'color': (0x4F, 0x81, 0xBD),  # Light blue #4F81BD
    },
}

# Table style colors (Grid Table 4 - black header)
TABLE_HEADER_BG = RGBColor(0x00, 0x00, 0x00)  # Black
TABLE_HEADER_TEXT = RGBColor(0xFF, 0xFF, 0xFF)  # White
TABLE_ALT_ROW_BG = RGBColor(0xF2, 0xF2, 0xF2)  # Light gray


# =============================================================================
# FORMATTING FUNCTIONS
# =============================================================================

def set_narrow_margins(doc):
    """Set narrow margins (1.27 cm / 0.5 inch) on all sections."""
    for section in doc.sections:
        section.top_margin = Cm(MARGIN_CM)
        section.bottom_margin = Cm(MARGIN_CM)
        section.left_margin = Cm(MARGIN_CM)
        section.right_margin = Cm(MARGIN_CM)


def apply_heading_styles(doc):
    """Apply custom formatting to heading styles."""
    for style_name, fmt in HEADING_STYLES.items():
        try:
            style = doc.styles[style_name]
            font = style.font
            font.size = Pt(fmt['size'])
            font.bold = fmt['bold']
            font.color.rgb = RGBColor(*fmt['color'])
        except KeyError:
            pass


def remove_trailing_paragraph_marks(doc):
    """Remove trailing ¶ symbols (newlines/spaces) from headings."""
    for para in doc.paragraphs:
        if para.style and para.style.name in ['Title', 'Heading 1', 'Heading 2', 'Heading 3']:
            # Strip trailing whitespace from each run
            for run in para.runs:
                if run.text:
                    run.text = run.text.rstrip('\n\r \t')


def is_placeholder_image(inline_shape):
    """
    Check if image is a gray placeholder (matplotlib toolbar icon).
    These are typically very small or have specific dimensions.
    """
    try:
        # Matplotlib toolbar icons are usually small (< 50 pixels)
        width_px = inline_shape.width.emu / 914400 * 96  # Convert EMU to pixels
        height_px = inline_shape.height.emu / 914400 * 96
        
        # Skip very small images (toolbar icons, buttons)
        if width_px < 100 and height_px < 100:
            return True
        
        # Skip if width/height ratio suggests an icon
        if width_px < 50 or height_px < 50:
            return True
            
        return False
    except:
        return False


def process_images(doc):
    """
    Process all images:
    - Remove placeholder/toolbar images
    - Lock aspect ratio
    - Scale to 70%
    - Center
    """
    processed = 0
    removed = 0
    
    # Process inline shapes (images)
    for para in doc.paragraphs:
        inline_shapes_to_remove = []
        
        for inline_shape in para._element.xpath('.//w:drawing'):
            # Get the inline element
            try:
                # Check dimensions to identify placeholders
                extent = inline_shape.xpath('.//wp:extent', namespaces={
                    'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
                })
                
                if extent:
                    cx = int(extent[0].get('cx', 0))
                    cy = int(extent[0].get('cy', 0))
                    
                    # Convert EMU to cm (914400 EMU = 1 inch, 1 inch = 2.54 cm)
                    width_cm = cx / 914400 * 2.54
                    height_cm = cy / 914400 * 2.54
                    
                    # Skip small images (toolbar icons, < 2cm)
                    if width_cm < 2 or height_cm < 2:
                        inline_shapes_to_remove.append(inline_shape)
                        removed += 1
                        continue
                    
                    # Scale to 70%
                    new_cx = int(cx * IMAGE_SCALE)
                    new_cy = int(cy * IMAGE_SCALE)
                    extent[0].set('cx', str(new_cx))
                    extent[0].set('cy', str(new_cy))
                    
                    # Also update the extent in a]graphic element if present
                    a_ext = inline_shape.xpath('.//a:ext', namespaces={
                        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'
                    })
                    for ext in a_ext:
                        if ext.get('cx'):
                            ext.set('cx', str(new_cx))
                        if ext.get('cy'):
                            ext.set('cy', str(new_cy))
                    
                    processed += 1
                    
            except Exception as e:
                pass
        
        # Remove placeholder images
        for shape in inline_shapes_to_remove:
            shape.getparent().remove(shape)
        
        # Center paragraph if it contains remaining images
        if para._element.xpath('.//w:drawing'):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return processed, removed


def set_cell_shading(cell, color):
    """Set background color for a table cell."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), f'{color.red:02X}{color.green:02X}{color.blue:02X}')
    cell._tc.get_or_add_tcPr().append(shading)


def format_tables(doc):
    """Apply Grid Table 4 style formatting to all tables."""
    formatted = 0
    
    for table in doc.tables:
        try:
            # Set table alignment to center
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Process rows
            for row_idx, row in enumerate(table.rows):
                for cell in row.cells:
                    # Header row (first row) - black background, white text
                    if row_idx == 0:
                        set_cell_shading(cell, TABLE_HEADER_BG)
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.font.color.rgb = TABLE_HEADER_TEXT
                                run.font.bold = True
                    # Alternating rows - light gray for even rows
                    elif row_idx % 2 == 0:
                        set_cell_shading(cell, TABLE_ALT_ROW_BG)
            
            # Add borders to all cells
            tbl = table._tbl
            tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
            
            tblBorders = OxmlElement('w:tblBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:color'), '000000')
                tblBorders.append(border)
            
            # Remove existing borders and add new
            existing_borders = tblPr.find(qn('w:tblBorders'))
            if existing_borders is not None:
                tblPr.remove(existing_borders)
            tblPr.append(tblBorders)
            
            if tbl.tblPr is None:
                tbl.insert(0, tblPr)
            
            formatted += 1
            
        except Exception as e:
            print(f"      [WARN] Table formatting error: {e}")
    
    return formatted


def format_docx(docx_path: Path) -> dict:
    """Apply all formatting to the docx file."""
    stats = {'images': 0, 'removed': 0, 'tables': 0}
    
    try:
        doc = Document(str(docx_path))
        
        # 1. Set narrow margins
        set_narrow_margins(doc)
        
        # 2. Apply heading styles
        apply_heading_styles(doc)
        
        # 3. Remove trailing ¶ from headings
        remove_trailing_paragraph_marks(doc)
        
        # 4. Process images (remove placeholders, scale 70%, center, lock ratio)
        stats['images'], stats['removed'] = process_images(doc)
        
        # 5. Format tables (Grid Table 4 style)
        stats['tables'] = format_tables(doc)
        
        # Save
        doc.save(str(docx_path))
        
        return stats
        
    except Exception as e:
        print(f"      [ERROR] Formatting failed: {e}")
        return stats


def convert_notebook(ipynb_path: Path, output_folder: Path) -> bool:
    """
    Convert ipynb → HTML → DOCX with formatting.
    """
    print(f"\n  Processing: {ipynb_path.name}")
    
    docx_path = output_folder / f"{ipynb_path.stem}.docx"
    
    try:
        # Step 1: Read notebook
        print("    [1/3] Reading notebook...")
        with open(ipynb_path, 'r', encoding='utf-8') as f:
            notebook = nbformat.read(f, as_version=4)
        
        # Step 2: Convert to HTML then DOCX
        print("    [2/3] Converting to Word...")
        html_exporter = HTMLExporter()
        html_exporter.template_name = 'basic'
        (html_content, resources) = html_exporter.from_notebook_node(notebook)
        
        # Write HTML to temp file
        with tempfile.NamedTemporaryFile(mode='w', suffix='.html', 
                                          delete=False, encoding='utf-8') as tmp:
            tmp.write(html_content)
            tmp_html_path = tmp.name
        
        try:
            pypandoc.convert_file(tmp_html_path, 'docx', outputfile=str(docx_path))
        finally:
            Path(tmp_html_path).unlink(missing_ok=True)
        
        # Step 3: Apply formatting
        print("    [3/3] Formatting document...")
        stats = format_docx(docx_path)
        print(f"      Margins: narrow | Images: {stats['images']} scaled, {stats['removed']} removed | Tables: {stats['tables']}")
        
        print(f"    [SUCCESS] → {docx_path.name}")
        return True
        
    except Exception as e:
        print(f"    [ERROR] Conversion failed: {e}")
        import traceback
        traceback.print_exc()
        return False


def main():
    print("=" * 60)
    print("    IPYNB → Formatted Word Converter")
    print("    (Tables, Images, Margins, Styles)")
    print("=" * 60)
    
    # Get script's folder
    script_folder = Path(__file__).parent.resolve()
    print(f"\nScanning: {script_folder}")
    
    # Find notebooks
    notebooks = [f for f in script_folder.glob("*.ipynb") 
                 if ".ipynb_checkpoints" not in str(f)]
    
    if not notebooks:
        print("\n[INFO] No .ipynb files found in this folder.")
        input("\nPress Enter to exit...")
        return
    
    print(f"Found {len(notebooks)} notebook(s)")
    
    # Create output folder
    output_folder = script_folder / "Word_Outputs"
    output_folder.mkdir(exist_ok=True)
    print(f"Output: {output_folder}")
    print("-" * 60)
    
    # Process each notebook
    success = 0
    failed = 0
    
    for nb in notebooks:
        if convert_notebook(nb, output_folder):
            success += 1
        else:
            failed += 1
    
    # Summary
    print("\n" + "=" * 60)
    print(f"COMPLETE! Success: {success} | Failed: {failed}")
    print("=" * 60)
    
    input("\nPress Enter to exit...")


if __name__ == "__main__":
    main()
