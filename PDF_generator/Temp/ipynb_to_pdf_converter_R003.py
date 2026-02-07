"""
IPYNB to Formatted PDF Converter (Windows)
==========================================
All-in-one converter that:
  1. Converts .ipynb to .docx (via pypandoc)
  2. Centers all images & locks aspect ratio
  3. Applies custom heading styles (Title, H1, H2, H3)
  4. Converts to PDF (via MS Word)

Usage:
  Just double-click. Processes all .ipynb files in script's folder.

Output:
  - Word_Outputs/ folder with formatted .docx files
  - PDF_Outputs/ folder with final .pdf files

Requirements:
  - Windows with MS Word installed
  - pip install pypandoc python-docx docx2pdf
"""

import sys
from pathlib import Path

# Check dependencies
missing = []
try:
    import pypandoc
except ImportError:
    missing.append("pypandoc")

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    missing.append("python-docx")

try:
    from docx2pdf import convert as convert_to_pdf
except ImportError:
    missing.append("docx2pdf")

if missing:
    print("=" * 50)
    print("ERROR: Missing packages!")
    print(f"Run: pip install {' '.join(missing)}")
    print("=" * 50)
    input("Press Enter to exit...")
    sys.exit(1)


# =============================================================================
# HEADING STYLE CONFIGURATION (edit these to match your preferences)
# =============================================================================
HEADING_STYLES = {
    'Title': {
        'size': 26,
        'bold': False,
        'color': (0x17, 0x36, 0x5D),  # Dark navy blue
    },
    'Heading 1': {
        'size': 14,
        'bold': True,
        'color': (0x36, 0x5F, 0x91),  # Medium blue
    },
    'Heading 2': {
        'size': 13,
        'bold': True,
        'color': (0x4F, 0x81, 0xBD),  # Lighter blue
    },
    'Heading 3': {
        'size': 11,
        'bold': True,
        'color': (0x4F, 0x81, 0xBD),  # Lighter blue
    },
}


def apply_heading_styles(doc):
    """Apply custom formatting to heading styles."""
    for style_name, formatting in HEADING_STYLES.items():
        try:
            style = doc.styles[style_name]
            font = style.font
            font.size = Pt(formatting['size'])
            font.bold = formatting['bold']
            font.color.rgb = RGBColor(*formatting['color'])
        except KeyError:
            print(f"    [WARN] Style '{style_name}' not found in document")


def center_all_images(doc):
    """Center all images in the document (replicates VBA script)."""
    centered_count = 0
    
    for paragraph in doc.paragraphs:
        # Check if paragraph contains inline shapes (images)
        if paragraph._element.xpath('.//w:drawing'):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            centered_count += 1
    
    # Also handle images via inline shapes in runs
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run._element.xpath('.//w:drawing'):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                centered_count += 1
                break  # Only count paragraph once
    
    return centered_count


def format_docx(docx_path: Path) -> bool:
    """Apply formatting: center images + heading styles."""
    try:
        doc = Document(str(docx_path))
        
        # Apply heading styles
        apply_heading_styles(doc)
        
        # Center images
        img_count = center_all_images(doc)
        
        # Save
        doc.save(str(docx_path))
        print(f"    [FORMAT] Styled headings, centered {img_count} image(s)")
        return True
        
    except Exception as e:
        print(f"    [ERROR] Formatting failed: {e}")
        return False


def convert_notebook(ipynb_path: Path, docx_folder: Path, pdf_folder: Path) -> bool:
    """Full pipeline: ipynb → formatted docx → pdf."""
    
    print(f"\n  Processing: {ipynb_path.name}")
    
    docx_path = docx_folder / f"{ipynb_path.stem}.docx"
    pdf_path = pdf_folder / f"{ipynb_path.stem}.pdf"
    
    # Step 1: Convert ipynb to docx
    print("    [1/3] Converting to Word...")
    try:
        pypandoc.convert_file(str(ipynb_path), 'docx', outputfile=str(docx_path))
    except Exception as e:
        print(f"    [ERROR] Pandoc conversion failed: {e}")
        return False
    
    # Step 2: Format the docx (images + headings)
    print("    [2/3] Formatting document...")
    format_docx(docx_path)
    
    # Step 3: Convert to PDF
    print("    [3/3] Converting to PDF...")
    try:
        convert_to_pdf(str(docx_path), str(pdf_path))
        print(f"    [SUCCESS] → {pdf_path.name}")
        return True
    except Exception as e:
        print(f"    [ERROR] PDF conversion failed: {e}")
        return False


def main():
    print("=" * 55)
    print("    IPYNB → Formatted DOCX → PDF Converter")
    print("=" * 55)
    
    # Get script's folder
    script_folder = Path(__file__).parent.resolve()
    print(f"\nScanning: {script_folder}")
    
    # Find notebooks
    notebooks = [f for f in script_folder.glob("*.ipynb") 
                 if ".ipynb_checkpoints" not in str(f)]
    
    if not notebooks:
        print("\n[INFO] No .ipynb files found in this folder.")
        input("\nPress Enter to exit...")
        return
    
    print(f"Found {len(notebooks)} notebook(s)")
    
    # Create output folders
    docx_folder = script_folder / "Word_Outputs"
    pdf_folder = script_folder / "PDF_Outputs"
    docx_folder.mkdir(exist_ok=True)
    pdf_folder.mkdir(exist_ok=True)
    
    print(f"\nWord output: {docx_folder}")
    print(f"PDF output:  {pdf_folder}")
    print("-" * 55)
    
    # Process each notebook
    success = 0
    failed = 0
    
    for nb in notebooks:
        if convert_notebook(nb, docx_folder, pdf_folder):
            success += 1
        else:
            failed += 1
    
    # Summary
    print("\n" + "=" * 55)
    print(f"COMPLETE! Success: {success} | Failed: {failed}")
    print("=" * 55)
    
    input("\nPress Enter to exit...")


if __name__ == "__main__":
    main()
