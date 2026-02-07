"""
IPYNB to Formatted Word Converter (Windows)
============================================
Converts notebooks to Word with:
  - LaTeX equations preserved (via pandoc)
  - Syntax highlighting for code
  - Grid Table 4 style for DataFrames
  - Images scaled to 70%, centered
  - Narrow margins

Pipeline: ipynb → Clean HTML → DOCX (via Pandoc) → Format

Usage:
  Double-click to run. Processes all .ipynb in script's folder.

Output:
  Word_Outputs/ folder with formatted .docx files

Requirements:
  - pip install nbconvert pypandoc python-docx nbformat
  - Pandoc installed (pypandoc.download_pandoc() if needed)
"""

import sys
import tempfile
import re
import base64
from pathlib import Path

# Check dependencies
missing = []
try:
    import nbconvert
    from nbconvert import HTMLExporter
    from nbconvert.preprocessors import Preprocessor
except ImportError:
    missing.append("nbconvert")

try:
    import pypandoc
except ImportError:
    missing.append("pypandoc")

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    missing.append("python-docx")

try:
    import nbformat
except ImportError:
    missing.append("nbformat")

if missing:
    print("=" * 50)
    print("ERROR: Missing packages!")
    print(f"Run: pip install {' '.join(missing)}")
    print("=" * 50)
    input("Press Enter to exit...")
    sys.exit(1)


# =============================================================================
# CONFIGURATION
# =============================================================================

IMAGE_SCALE = 0.70  # 70%

# Heading styles
HEADING_STYLES = {
    'Title': {'size': 26, 'bold': False, 'color': (0x17, 0x36, 0x5D)},
    'Heading 1': {'size': 14, 'bold': True, 'color': (0x36, 0x5F, 0x91)},
    'Heading 2': {'size': 13, 'bold': True, 'color': (0x4F, 0x81, 0xBD)},
    'Heading 3': {'size': 11, 'bold': True, 'color': (0x4F, 0x81, 0xBD)},
}

# Patterns to remove from output (matplotlib artifacts, etc.)
UNWANTED_OUTPUT_PATTERNS = [
    r'<Axes:.*?>',
    r'<matplotlib\..*?>',
    r'<Figure.*?>',
    r'<AxesSubplot:.*?>',
    r'^<div>$',
    r'^</div>$',
    r'^\s*<div>\s*$',
    r'^\s*</div>\s*$',
]


# =============================================================================
# NOTEBOOK CLEANING
# =============================================================================

def clean_cell_outputs(notebook):
    """
    Clean notebook cell outputs to remove unwanted artifacts.
    - Remove matplotlib/axes text outputs
    - Keep images, tables, and meaningful text
    """
    for cell in notebook.cells:
        if cell.cell_type == 'code' and 'outputs' in cell:
            cleaned_outputs = []
            for output in cell.outputs:
                # Skip stream outputs that match unwanted patterns
                if output.get('output_type') == 'stream':
                    text = output.get('text', '')
                    skip = False
                    for pattern in UNWANTED_OUTPUT_PATTERNS:
                        if re.search(pattern, text, re.MULTILINE):
                            skip = True
                            break
                    if not skip:
                        cleaned_outputs.append(output)
                
                # For execute_result or display_data
                elif output.get('output_type') in ['execute_result', 'display_data']:
                    data = output.get('data', {})
                    
                    # Check text/plain for unwanted patterns
                    text_plain = data.get('text/plain', '')
                    if isinstance(text_plain, list):
                        text_plain = ''.join(text_plain)
                    
                    skip_text = False
                    for pattern in UNWANTED_OUTPUT_PATTERNS:
                        if re.search(pattern, text_plain):
                            skip_text = True
                            break
                    
                    # If has image, keep it regardless of text
                    if 'image/png' in data or 'image/jpeg' in data:
                        # Remove the text/plain if it's unwanted
                        if skip_text and 'text/plain' in data:
                            del data['text/plain']
                        cleaned_outputs.append(output)
                    # If has HTML (like DataFrames), keep it
                    elif 'text/html' in data:
                        cleaned_outputs.append(output)
                    # Only keep text/plain if it's not unwanted
                    elif not skip_text and text_plain.strip():
                        cleaned_outputs.append(output)
                
                else:
                    # Keep other output types (errors, etc.)
                    cleaned_outputs.append(output)
            
            cell.outputs = cleaned_outputs
    
    return notebook


# =============================================================================
# CUSTOM HTML TEMPLATE FOR BETTER CONVERSION
# =============================================================================

CUSTOM_HTML_TEMPLATE = """
{%- extends 'basic/index.html.j2' -%}

{% block header %}
<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
/* Code highlighting */
.highlight { background: #f8f8f8; padding: 10px; margin: 5px 0; }
.highlight .c1 { color: #408080; font-style: italic; } /* Comment */
.highlight .k { color: #008000; font-weight: bold; } /* Keyword */
.highlight .kn { color: #008000; font-weight: bold; } /* Keyword.Namespace */
.highlight .n { color: #000000; } /* Name */
.highlight .nn { color: #0000ff; font-weight: bold; } /* Name.Namespace */
.highlight .nf { color: #0000ff; } /* Name.Function */
.highlight .nb { color: #008000; } /* Name.Builtin */
.highlight .s1, .highlight .s2 { color: #ba2121; } /* String */
.highlight .mi { color: #666666; } /* Number */
.highlight .o { color: #666666; } /* Operator */
.highlight .p { color: #000000; } /* Punctuation */

/* Tables */
table { border-collapse: collapse; margin: 10px 0; }
th, td { border: 1px solid black; padding: 5px; text-align: center; }
th { background-color: black; color: white; font-weight: bold; }
tr:nth-child(even) { background-color: #f2f2f2; }

/* Output */
.output_text pre { background: #f8f8f8; padding: 5px; }
</style>
</head>
<body>
{%- endblock header %}

{% block footer %}
</body>
</html>
{%- endblock footer %}
"""


# =============================================================================
# FORMATTING FUNCTIONS
# =============================================================================

def set_narrow_margins(doc):
    """Set margins to Narrow (0.5 inches)."""
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)


def apply_heading_styles(doc):
    """Apply custom heading styles."""
    for style_name, fmt in HEADING_STYLES.items():
        try:
            style = doc.styles[style_name]
            font = style.font
            font.size = Pt(fmt['size'])
            font.bold = fmt['bold']
            font.color.rgb = RGBColor(*fmt['color'])
        except KeyError:
            pass


def remove_trailing_paragraph_marks(doc):
    """Remove trailing newlines from headings."""
    for para in doc.paragraphs:
        if para.style and para.style.name in ['Title', 'Heading 1', 'Heading 2', 'Heading 3']:
            for run in para.runs:
                if run.text:
                    run.text = run.text.rstrip('\n\r \t')


def set_cell_shading(cell, hex_color):
    """Set background shading for a table cell."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), hex_color)
    shading.set(qn('w:val'), 'clear')
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn('w:shd'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(shading)


def format_tables(doc):
    """
    Format tables to Grid Table 4 style.
    - Cambria font, 11pt
    - Header row: black background, white bold text
    - Banded rows (alternating gray)
    - AutoFit, centered
    """
    formatted = 0
    
    for table in doc.tables:
        try:
            # Try built-in style first
            try:
                table.style = 'Grid Table 4'
            except:
                pass
            
            table.autofit = True
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Set table width to 100%
            tbl = table._tbl
            tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
            
            tblW = tblPr.find(qn('w:tblW'))
            if tblW is None:
                tblW = OxmlElement('w:tblW')
                tblPr.append(tblW)
            tblW.set(qn('w:w'), '5000')
            tblW.set(qn('w:type'), 'pct')
            
            if tbl.tblPr is None:
                tbl.insert(0, tblPr)
            
            # Process each row
            for row_idx, row in enumerate(table.rows):
                # Center row
                row_tr = row._tr
                trPr = row_tr.get_or_add_trPr()
                jc = trPr.find(qn('w:jc'))
                if jc is None:
                    jc = OxmlElement('w:jc')
                    trPr.append(jc)
                jc.set(qn('w:val'), 'center')
                
                for cell in row.cells:
                    # Set font
                    for para in cell.paragraphs:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in para.runs:
                            run.font.name = 'Cambria'
                            run.font.size = Pt(11)
                    
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    
                    # Header row
                    if row_idx == 0:
                        set_cell_shading(cell, '000000')
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                                run.font.bold = True
                    # Banded rows
                    elif row_idx % 2 == 0:
                        set_cell_shading(cell, 'F2F2F2')
            
            # Add borders
            tblBorders = OxmlElement('w:tblBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:color'), '000000')
                tblBorders.append(border)
            
            existing_borders = tblPr.find(qn('w:tblBorders'))
            if existing_borders is not None:
                tblPr.remove(existing_borders)
            tblPr.append(tblBorders)
            
            formatted += 1
            
        except Exception as e:
            print(f"      [WARN] Table error: {e}")
    
    return formatted


def process_images(doc):
    """
    Process all images.
    - Scale to 70%
    - Center paragraph
    """
    processed = 0
    
    for shape in doc.inline_shapes:
        try:
            inline = shape._inline
            
            # Find extent element
            extent = inline.find(qn('wp:extent'))
            if extent is not None:
                cx = int(extent.get('cx'))
                cy = int(extent.get('cy'))
                
                # Scale to 70%
                new_cx = int(cx * IMAGE_SCALE)
                new_cy = int(cy * IMAGE_SCALE)
                
                extent.set('cx', str(new_cx))
                extent.set('cy', str(new_cy))
                
                # Update graphic extent
                for ext in inline.iter(qn('a:ext')):
                    if ext.get('cx') and ext.get('cy'):
                        ext.set('cx', str(new_cx))
                        ext.set('cy', str(new_cy))
            
            # Center the paragraph
            parent = inline.getparent()
            while parent is not None:
                if parent.tag == qn('w:p'):
                    pPr = parent.find(qn('w:pPr'))
                    if pPr is None:
                        pPr = OxmlElement('w:pPr')
                        parent.insert(0, pPr)
                    
                    jc = pPr.find(qn('w:jc'))
                    if jc is None:
                        jc = OxmlElement('w:jc')
                        pPr.append(jc)
                    jc.set(qn('w:val'), 'center')
                    break
                parent = parent.getparent()
            
            processed += 1
            
        except Exception:
            pass
    
    return processed


def remove_unwanted_paragraphs(doc):
    """Remove paragraphs containing only unwanted artifacts."""
    removed = 0
    paragraphs_to_remove = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            for pattern in UNWANTED_OUTPUT_PATTERNS:
                if re.fullmatch(pattern, text):
                    paragraphs_to_remove.append(para)
                    break
    
    for para in paragraphs_to_remove:
        try:
            p = para._element
            p.getparent().remove(p)
            removed += 1
        except:
            pass
    
    return removed


def format_docx(docx_path: Path) -> dict:
    """Apply all formatting to the docx file."""
    stats = {'images': 0, 'tables': 0, 'removed': 0}
    
    try:
        doc = Document(str(docx_path))
        
        # 1. Set narrow margins
        set_narrow_margins(doc)
        
        # 2. Apply heading styles
        apply_heading_styles(doc)
        
        # 3. Remove trailing marks from headings
        remove_trailing_paragraph_marks(doc)
        
        # 4. Remove unwanted paragraphs
        stats['removed'] = remove_unwanted_paragraphs(doc)
        
        # 5. Format tables
        stats['tables'] = format_tables(doc)
        
        # 6. Process images
        stats['images'] = process_images(doc)
        
        # Save
        doc.save(str(docx_path))
        
        return stats
        
    except Exception as e:
        print(f"      [ERROR] Formatting failed: {e}")
        import traceback
        traceback.print_exc()
        return stats


def convert_notebook(ipynb_path: Path, output_folder: Path) -> bool:
    """Convert ipynb → HTML → DOCX with formatting."""
    print(f"\n  Processing: {ipynb_path.name}")
    
    docx_path = output_folder / f"{ipynb_path.stem}.docx"
    
    try:
        # Step 1: Read and clean notebook
        print("    [1/4] Reading notebook...")
        with open(ipynb_path, 'r', encoding='utf-8') as f:
            notebook = nbformat.read(f, as_version=4)
        
        # Clean unwanted outputs
        notebook = clean_cell_outputs(notebook)
        
        # Step 2: Convert to HTML
        print("    [2/4] Converting to HTML...")
        html_exporter = HTMLExporter()
        html_exporter.template_name = 'classic'  # Better syntax highlighting
        (html_content, resources) = html_exporter.from_notebook_node(notebook)
        
        # Create temp directory for resources
        tmp_dir = Path(tempfile.mkdtemp(prefix="ipynb2docx_"))
        tmp_html_path = tmp_dir / "notebook.html"
        
        # Write HTML
        tmp_html_path.write_text(html_content, encoding='utf-8')
        
        # Write any extracted images
        for fname, data in (resources.get('outputs') or {}).items():
            out_path = tmp_dir / fname
            out_path.parent.mkdir(parents=True, exist_ok=True)
            if isinstance(data, bytes):
                out_path.write_bytes(data)
            else:
                out_path.write_text(str(data), encoding='utf-8')
        
        # Step 3: Convert HTML to DOCX via Pandoc
        print("    [3/4] Converting to Word...")
        try:
            pypandoc.convert_file(
                str(tmp_html_path),
                'docx',
                outputfile=str(docx_path),
                extra_args=[
                    '--from=html',
                    '--standalone',
                    f'--resource-path={str(tmp_dir)}',
                    '--wrap=none',
                ]
            )
        finally:
            # Cleanup
            try:
                for p in sorted(tmp_dir.rglob("*"), reverse=True):
                    if p.is_file():
                        p.unlink(missing_ok=True)
                    elif p.is_dir():
                        p.rmdir()
                tmp_dir.rmdir()
            except:
                pass
        
        # Step 4: Apply formatting
        print("    [4/4] Formatting document...")
        stats = format_docx(docx_path)
        print(f"      Tables: {stats['tables']} | Images: {stats['images']} | Cleaned: {stats['removed']}")
        
        print(f"    [SUCCESS] → {docx_path.name}")
        return True
        
    except Exception as e:
        print(f"    [ERROR] Conversion failed: {e}")
        import traceback
        traceback.print_exc()
        return False


def main():
    print("=" * 60)
    print("    IPYNB → Formatted Word Converter")
    print("    (Equations + Highlighting + Tables + Images)")
    print("=" * 60)
    
    script_folder = Path(__file__).parent.resolve()
    print(f"\nScanning: {script_folder}")
    
    notebooks = [f for f in script_folder.glob("*.ipynb") 
                 if ".ipynb_checkpoints" not in str(f)]
    
    if not notebooks:
        print("\n[INFO] No .ipynb files found in this folder.")
        input("\nPress Enter to exit...")
        return
    
    print(f"Found {len(notebooks)} notebook(s)")
    
    output_folder = script_folder / "Word_Outputs"
    output_folder.mkdir(exist_ok=True)
    print(f"Output: {output_folder}")
    print("-" * 60)
    
    success = 0
    failed = 0
    
    for nb in notebooks:
        if convert_notebook(nb, output_folder):
            success += 1
        else:
            failed += 1
    
    print("\n" + "=" * 60)
    print(f"COMPLETE! Success: {success} | Failed: {failed}")
    print("=" * 60)
    
    input("\nPress Enter to exit...")


if __name__ == "__main__":
    main()
