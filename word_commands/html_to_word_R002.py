"""
HTML to Formatted Word Converter
================================
Takes custom notebook HTML files (from notebook_to_html) and converts them
to properly formatted Word documents via Pandoc + python-docx post-processing.

Features:
  - Syntax highlighting colors preserved (keywords, strings, numbers, comments)
  - Code input vs output visually distinct
  - Narrow margins (0.5") - no wasted space
  - Compact paragraph spacing - no blank pages
  - Images scaled to fit page
  - Tables with dark headers + zebra stripes
  - Matplotlib artifacts removed
  - Heading styles with colors

Usage:
    # Auto-convert all .html in script's folder (double-click to run)
    python html_to_word.py

    # Single file
    python html_to_word.py notebook.html
    python html_to_word.py notebook.html -o output.docx

    # Batch convert a directory
    python html_to_word.py --batch /path/to/html/
    python html_to_word.py --batch /path/to/html/ -o /path/to/output/

Requirements:
    - pip install python-docx
    - Pandoc installed (https://pandoc.org)
"""

import sys
import re
import os
import copy
import subprocess
import shutil
import argparse
import tempfile
from pathlib import Path
from html import unescape

# -- Dependencies --
if not shutil.which('pandoc'):
    print("ERROR: pandoc not found. Install from https://pandoc.org")
    sys.exit(1)

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Emu, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: pip install python-docx")
    sys.exit(1)


# ======================================================================
# CONFIGURATION
# ======================================================================

MARGIN = Inches(0.5)
MAX_IMAGE_WIDTH = Inches(6.5)
MAX_IMAGE_HEIGHT = Inches(4.5)
IMAGE_SCALE = 0.80

CODE_FONT = 'Consolas'
CODE_INPUT_SIZE = Pt(9)
CODE_OUTPUT_SIZE = Pt(8.5)
CODE_OUTPUT_COLOR = RGBColor(0x55, 0x55, 0x55)  # Grey for output

HEADING_STYLES = {
    'Title':     {'size': 26, 'bold': True,  'color': (0x17, 0x36, 0x5D)},
    'Heading 1': {'size': 16, 'bold': True,  'color': (0x36, 0x5F, 0x91)},
    'Heading 2': {'size': 14, 'bold': True,  'color': (0x4F, 0x81, 0xBD)},
    'Heading 3': {'size': 12, 'bold': True,  'color': (0x4F, 0x81, 0xBD)},
    'Heading 4': {'size': 11, 'bold': True,  'color': (0x4F, 0x81, 0xBD)},
}

STYLE_SPACING = {
    'Title':           {'before': 6,  'after': 3,  'line': 1.0},
    'Heading 1':       {'before': 12, 'after': 3,  'line': 1.0},
    'Heading 2':       {'before': 10, 'after': 2,  'line': 1.0},
    'Heading 3':       {'before': 8,  'after': 2,  'line': 1.0},
    'Heading 4':       {'before': 8,  'after': 2,  'line': 1.0},
    'Body Text':       {'before': 2,  'after': 2,  'line': 1.0},
    'First Paragraph': {'before': 2,  'after': 2,  'line': 1.0},
    'Compact':         {'before': 1,  'after': 1,  'line': 1.0},
    'Source Code':     {'before': 1,  'after': 1,  'line': 1.0},
    'Block Text':      {'before': 2,  'after': 2,  'line': 1.0},
    'Normal':          {'before': 2,  'after': 2,  'line': 1.0},
}

UNWANTED_PATTERNS = [
    r'<Axes:.*?>',
    r'<AxesSubplot:.*?>',
    r'<matplotlib\..*?>',
    r'<Figure.*?>',
    r'<mpl_toolkits\..*?>',
]


# ======================================================================
# SYNTAX COLOR EXTRACTION FROM HTML
# ======================================================================

def extract_color_map(css_text):
    """Extract CSS class -> color mapping from <style> block."""
    color_map = {}
    SKIP = {'highlight', 'code', 'output', 'cell', 'input',
            'notebook', 'container', 'markdown'}

    for m in re.finditer(r'([^{}]+)\{([^}]*(?<![a-z-])color\s*:[^}]+)\}', css_text):
        selectors, props = m.group(1), m.group(2)
        c = re.search(r'(?<![a-z-])color\s*:\s*(#[0-9a-fA-F]{3,8})', props)
        if not c:
            continue
        hex_color = c.group(1)
        if len(hex_color) == 4:
            hex_color = '#' + hex_color[1]*2 + hex_color[2]*2 + hex_color[3]*2
        bold = bool(re.search(r'font-weight\s*:\s*(bold|[6-9]00)', props))
        italic = bool(re.search(r'font-style\s*:\s*italic', props))

        for cls_m in re.finditer(r'\.([a-zA-Z_][\w]*)', selectors):
            cls = cls_m.group(1)
            if cls not in SKIP:
                color_map[cls] = {'color': hex_color, 'bold': bold, 'italic': italic}

    return color_map


def parse_spans(block_html, color_map):
    """Parse code block HTML into [(text, hex_color, bold, italic), ...]."""
    fragments = []
    pos = 0

    for m in re.finditer(r'<span class="(\w+)">(.*?)</span>', block_html, re.DOTALL):
        before = block_html[pos:m.start()]
        before = re.sub(r'<[^>]+>', '', before)
        before = unescape(before)
        if before:
            fragments.append((before, None, False, False))

        cls = m.group(1)
        text = re.sub(r'<[^>]+>', '', m.group(2))
        text = unescape(text)
        info = color_map.get(cls)
        color = info['color'] if info else None
        bold = info['bold'] if info else False
        italic = info['italic'] if info else False
        if text:
            fragments.append((text, color, bold, italic))
        pos = m.end()

    after = block_html[pos:]
    after = re.sub(r'<[^>]+>', '', after)
    after = unescape(after)
    if after:
        fragments.append((after, None, False, False))

    return fragments


def extract_ordered_blocks(html_text, color_map):
    """Extract ALL code blocks (input + output) from HTML in document order."""
    blocks = []
    pattern = r'<div class="(code-input|code-output)">(.*?)</div>'

    for m in re.finditer(pattern, html_text, re.DOTALL):
        block_type = 'input' if m.group(1) == 'code-input' else 'output'
        content = m.group(2)

        if block_type == 'input':
            inner = re.sub(r'<pre>\s*<code[^>]*>(.*?)</code>\s*</pre>',
                           r'\1', content, flags=re.DOTALL)
            fragments = parse_spans(inner, color_map)
        else:
            inner = re.sub(r'</?pre[^>]*>', '', content)
            inner = re.sub(r'<[^>]+>', '', inner)
            text = unescape(inner)
            fragments = [(text, None, False, False)]

        plain_text = ''.join(t for t, c, b, i in fragments).strip()
        blocks.append({
            'type': block_type,
            'fragments': fragments,
            'text': plain_text,
        })

    return blocks


def hex_to_rgb(hex_color):
    """Convert #RRGGBB to RGBColor."""
    h = hex_color.lstrip('#')
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def rebuild_para_with_colors(para, fragments, is_input):
    """Clear a paragraph's runs and rebuild with colored runs."""
    # Remove ALL existing runs and hyperlinks
    for child in list(para._element):
        if child.tag in (qn('w:r'), qn('w:hyperlink')):
            para._element.remove(child)

    font_size = CODE_INPUT_SIZE if is_input else CODE_OUTPUT_SIZE

    for text, color_hex, bold, italic in fragments:
        if not text:
            continue

        parts = text.split('\n')
        for i, part in enumerate(parts):
            if i > 0:
                # Line break run with font properties
                br_run = OxmlElement('w:r')
                rPr = OxmlElement('w:rPr')
                rFonts = OxmlElement('w:rFonts')
                rFonts.set(qn('w:ascii'), CODE_FONT)
                rFonts.set(qn('w:hAnsi'), CODE_FONT)
                rPr.append(rFonts)
                sz = OxmlElement('w:sz')
                sz.set(qn('w:val'), str(int(font_size.pt * 2)))
                rPr.append(sz)
                br_run.append(rPr)
                br_run.append(OxmlElement('w:br'))
                para._element.append(br_run)

            if part:
                run = para.add_run(part)
                run.font.name = CODE_FONT
                run.font.size = font_size

                if is_input and color_hex:
                    run.font.color.rgb = hex_to_rgb(color_hex)
                elif not is_input:
                    run.font.color.rgb = CODE_OUTPUT_COLOR

                if bold:
                    run.font.bold = True
                if italic:
                    run.font.italic = True


def normalize(text):
    """Normalize text for comparison."""
    return re.sub(r'\s+', ' ', text.strip())


def apply_syntax_colors(doc, html_text):
    """
    Post-process: apply syntax highlighting colors from HTML to docx.
    Matches Source Code paragraphs 1:1 with HTML code blocks.
    """
    css_match = re.search(r'<style>(.*?)</style>', html_text, re.DOTALL)
    css_text = css_match.group(1) if css_match else ""
    color_map = extract_color_map(css_text)

    if not color_map:
        return 0, 0

    blocks = extract_ordered_blocks(html_text, color_map)
    if not blocks:
        return 0, 0

    sc_paras = [p for p in doc.paragraphs
                if p.style and p.style.name == 'Source Code']

    colored_inputs = 0
    styled_outputs = 0
    block_idx = 0

    for para in sc_paras:
        if block_idx >= len(blocks):
            break

        para_text = para.text.strip()
        if not para_text:
            block_idx += 1
            continue

        block = blocks[block_idx]

        if normalize(para_text) == normalize(block['text']):
            if block['type'] == 'input':
                rebuild_para_with_colors(para, block['fragments'], is_input=True)
                colored_inputs += 1
            else:
                rebuild_para_with_colors(para, block['fragments'], is_input=False)
                styled_outputs += 1
            block_idx += 1
        else:
            # Lookahead to find match
            found = False
            for j in range(block_idx, min(block_idx + 5, len(blocks))):
                if normalize(para_text) == normalize(blocks[j]['text']):
                    block = blocks[j]
                    if block['type'] == 'input':
                        rebuild_para_with_colors(para, block['fragments'], is_input=True)
                        colored_inputs += 1
                    else:
                        rebuild_para_with_colors(para, block['fragments'], is_input=False)
                        styled_outputs += 1
                    block_idx = j + 1
                    found = True
                    break
            if not found:
                block_idx += 1

    return colored_inputs, styled_outputs


# ======================================================================
# DOCUMENT FORMATTING FUNCTIONS
# ======================================================================

def set_margins(doc):
    for section in doc.sections:
        section.top_margin = MARGIN
        section.bottom_margin = MARGIN
        section.left_margin = MARGIN
        section.right_margin = MARGIN
        if section.page_width is None:
            section.page_width = Inches(8.27)
            section.page_height = Inches(11.69)


def fix_style_spacing(doc):
    for style_name, spacing in STYLE_SPACING.items():
        try:
            style = doc.styles[style_name]
            pf = style.paragraph_format
            pf.space_before = Pt(spacing['before'])
            pf.space_after = Pt(spacing['after'])
            if 'line' in spacing:
                pf.line_spacing = spacing['line']
        except KeyError:
            pass


def apply_heading_styles(doc):
    for style_name, fmt in HEADING_STYLES.items():
        try:
            style = doc.styles[style_name]
            font = style.font
            font.size = Pt(fmt['size'])
            font.bold = fmt['bold']
            font.color.rgb = RGBColor(*fmt['color'])
        except KeyError:
            pass


def remove_duplicate_title(doc):
    if len(doc.paragraphs) < 2:
        return 0
    first, second = doc.paragraphs[0], doc.paragraphs[1]
    if (first.style and second.style and
        first.text.strip().rstrip('\u00b6') == second.text.strip().rstrip('\u00b6') and
        'Title' in first.style.name):
        first._element.getparent().remove(first._element)
        return 1
    return 0


def remove_unwanted_paragraphs(doc):
    to_remove = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            for pattern in UNWANTED_PATTERNS:
                if re.fullmatch(pattern, text):
                    to_remove.append(para)
                    break
    for para in to_remove:
        try:
            para._element.getparent().remove(para._element)
        except:
            pass
    return len(to_remove)


def clean_heading_anchors(doc):
    for para in doc.paragraphs:
        if para.style and 'Heading' in para.style.name:
            for run in para.runs:
                if run.text and run.text.endswith('\u00b6'):
                    run.text = run.text.rstrip('\u00b6').rstrip()
            for hl in para._element.findall(qn('w:hyperlink')):
                hl_text = ''.join(r.text or '' for r in hl.findall('.//' + qn('w:t')))
                if hl_text.strip() == '\u00b6':
                    para._element.remove(hl)


def process_images(doc):
    processed = 0
    for shape in doc.inline_shapes:
        try:
            inline = shape._inline
            extent = inline.find(qn('wp:extent'))
            if extent is None:
                continue
            cx, cy = int(extent.get('cx')), int(extent.get('cy'))
            max_w, max_h = int(MAX_IMAGE_WIDTH), int(MAX_IMAGE_HEIGHT)
            new_cx, new_cy = cx, cy
            if new_cx > max_w:
                ratio = max_w / new_cx
                new_cx, new_cy = max_w, int(new_cy * ratio)
            if new_cy > max_h:
                ratio = max_h / new_cy
                new_cy, new_cx = max_h, int(new_cx * ratio)
            new_cx = int(new_cx * IMAGE_SCALE)
            new_cy = int(new_cy * IMAGE_SCALE)
            extent.set('cx', str(new_cx))
            extent.set('cy', str(new_cy))
            for ext in inline.iter(qn('a:ext')):
                if ext.get('cx') and ext.get('cy'):
                    ext.set('cx', str(new_cx))
                    ext.set('cy', str(new_cy))
            parent = inline.getparent()
            while parent is not None:
                if parent.tag == qn('w:p'):
                    pPr = parent.find(qn('w:pPr'))
                    if pPr is None:
                        pPr = OxmlElement('w:pPr')
                        parent.insert(0, pPr)
                    jc = pPr.find(qn('w:jc'))
                    if jc is None:
                        jc = OxmlElement('w:jc')
                        pPr.append(jc)
                    jc.set(qn('w:val'), 'center')
                    break
                parent = parent.getparent()
            processed += 1
        except Exception:
            pass
    return processed


def set_cell_shading(cell, hex_color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), hex_color)
    shading.set(qn('w:val'), 'clear')
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn('w:shd'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(shading)


def set_cell_borders(cell, color='000000', size='4'):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement('w:{}'.format(edge))
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), size)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    existing = tcPr.find(qn('w:tcBorders'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(tcBorders)


def format_tables(doc):
    """Format tables: full page width, black header, bold index column, zebra stripes."""
    formatted = 0
    # Page width minus margins = usable width
    usable_width = Inches(7.27)  # 8.27 - 2*0.5

    for table in doc.tables:
        try:
            tblPr = table._tbl.tblPr
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                table._tbl.insert(0, tblPr)

            # Set table to full page width
            tblW = tblPr.find(qn('w:tblW'))
            if tblW is None:
                tblW = OxmlElement('w:tblW')
                tblPr.append(tblW)
            tblW.set(qn('w:type'), 'dxa')
            tblW.set(qn('w:w'), str(int(usable_width / Emu(635))))  # convert to DXA

            # Center table
            jc = tblPr.find(qn('w:jc'))
            if jc is None:
                jc = OxmlElement('w:jc')
                tblPr.append(jc)
            jc.set(qn('w:val'), 'center')

            num_cols = len(table.columns)

            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    set_cell_borders(cell, '999999', '4')

                    if row_idx == 0:
                        # Header row: black background, white bold text
                        set_cell_shading(cell, '000000')
                        for para in cell.paragraphs:
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in para.runs:
                                run.font.bold = True
                                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                                run.font.size = Pt(10)
                    else:
                        # Data rows: alternating grey/white
                        set_cell_shading(cell, 'E8E8E8' if row_idx % 2 == 0 else 'FFFFFF')

                        for para in cell.paragraphs:
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in para.runs:
                                run.font.size = Pt(10)

                                # Bold first column (index)
                                if col_idx == 0:
                                    run.font.bold = True

                    # Compact cell spacing
                    for para in cell.paragraphs:
                        para.paragraph_format.space_before = Pt(2)
                        para.paragraph_format.space_after = Pt(2)

            # Table-level borders
            tblBorders = OxmlElement('w:tblBorders')
            for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement('w:{}'.format(edge))
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '999999')
                tblBorders.append(border)
            existing = tblPr.find(qn('w:tblBorders'))
            if existing is not None:
                tblPr.remove(existing)
            tblPr.append(tblBorders)
            formatted += 1
        except Exception as e:
            print("  [WARN] Table format error: {}".format(e))
    return formatted


def split_large_code_paragraphs(doc):
    """
    Split large Source Code paragraphs (with many w:br line breaks) into
    separate paragraphs — one per line. This prevents blank pages caused by
    Word refusing to split a single paragraph across pages.
    """
    MAX_BREAKS = 6  # Split if more than 6 line breaks in one paragraph
    split_count = 0

    # Collect paragraphs to split (can't modify while iterating)
    to_split = []
    for para in doc.paragraphs:
        if para.style and para.style.name == 'Source Code':
            breaks = para._element.findall('.//' + qn('w:br'))
            if len(breaks) > MAX_BREAKS:
                to_split.append(para)

    for para in to_split:
        try:
            parent = para._element.getparent()
            index = list(parent).index(para._element)
            style_name = para.style.name

            # Collect all text lines by walking runs
            lines = []
            current_line_runs = []

            for child in para._element:
                if child.tag == qn('w:pPr'):
                    continue
                if child.tag == qn('w:r'):
                    has_br = child.find(qn('w:br')) is not None
                    if has_br:
                        # Save current line, start new one
                        # Get text from this run before the break
                        for t_elem in child.findall(qn('w:t')):
                            if t_elem.text:
                                rPr = child.find(qn('w:rPr'))
                                current_line_runs.append((t_elem.text, rPr))
                        lines.append(current_line_runs)
                        current_line_runs = []
                    else:
                        for t_elem in child.findall(qn('w:t')):
                            if t_elem.text:
                                rPr = child.find(qn('w:rPr'))
                                current_line_runs.append((t_elem.text, rPr))

            if current_line_runs:
                lines.append(current_line_runs)

            if len(lines) <= 1:
                continue

            # Remove the original paragraph
            parent.remove(para._element)

            # Create new paragraphs for each line
            for line_idx, line_runs in enumerate(lines):
                text = ''.join(t for t, rpr in line_runs).strip()
                if not text and line_idx > 0:
                    continue  # Skip empty lines (but keep first)

                new_p = OxmlElement('w:p')
                # Copy paragraph properties with style
                pPr = OxmlElement('w:pPr')
                pStyle = OxmlElement('w:pStyle')
                pStyle.set(qn('w:val'), 'SourceCode')
                pPr.append(pStyle)
                new_p.append(pPr)

                # Add runs with their formatting
                for run_text, run_rPr in line_runs:
                    new_r = OxmlElement('w:r')
                    if run_rPr is not None:
                        new_r.append(copy.deepcopy(run_rPr))
                    new_t = OxmlElement('w:t')
                    new_t.set(qn('xml:space'), 'preserve')
                    new_t.text = run_text
                    new_r.append(new_t)
                    new_p.append(new_r)

                parent.insert(index, new_p)
                index += 1

            split_count += 1

        except Exception as e:
            pass  # Skip problematic paragraphs

    return split_count


def reduce_paragraph_spacing(doc):
    fixed = 0
    max_before, max_after = Pt(14), Pt(6)
    for para in doc.paragraphs:
        pf = para.paragraph_format
        changed = False
        if pf.space_before and pf.space_before > max_before:
            pf.space_before = max_before
            changed = True
        if pf.space_after and pf.space_after > max_after:
            pf.space_after = max_after
            changed = True
        if changed:
            fixed += 1
    return fixed


# ======================================================================
# MAIN CONVERSION
# ======================================================================

def convert_html_to_docx(html_path, output_path):
    """
    Convert HTML -> DOCX via Pandoc, then post-process for formatting + syntax colors.
    """
    html_path = Path(html_path)
    output_path = Path(output_path)
    print("\n  Processing: {}".format(html_path.name))

    try:
        # Read original HTML
        with open(html_path, 'r', encoding='utf-8') as f:
            html_text = f.read()

        # Step 1: Pandoc
        print("    [1/4] Converting HTML -> DOCX via Pandoc...")
        result = subprocess.run([
            'pandoc', '--from=html', '--to=docx', '--standalone',
            '--resource-path={}'.format(str(html_path.parent)),
            '--wrap=none',
            '-o', str(output_path),
            str(html_path),
        ], capture_output=True, text=True)
        if result.returncode != 0:
            print("    [ERROR] Pandoc: {}".format(result.stderr))
            return None

        # Step 2: Syntax colors
        print("    [2/4] Applying syntax highlighting colors...")
        doc = Document(str(output_path))
        colored, styled = apply_syntax_colors(doc, html_text)

        # Step 3: Spacing & layout
        print("    [3/4] Fixing spacing & layout...")
        set_margins(doc)
        fix_style_spacing(doc)
        apply_heading_styles(doc)
        dup = remove_duplicate_title(doc)
        artifacts = remove_unwanted_paragraphs(doc)
        clean_heading_anchors(doc)
        splits = split_large_code_paragraphs(doc)
        spacing_fixed = reduce_paragraph_spacing(doc)

        # Step 4: Tables & images
        print("    [4/4] Formatting tables & images...")
        tables = format_tables(doc)
        images = process_images(doc)

        doc.save(str(output_path))

        print("      Code colored: {} input + {} output".format(colored, styled))
        print("      Tables: {} | Images: {} | Artifacts: {} | Splits: {}".format(tables, images, artifacts, splits))
        print("    > Generated: {}".format(output_path))
        return str(output_path)

    except Exception as e:
        print("    x Failed: {}".format(e))
        import traceback
        traceback.print_exc()
        return None


def batch_convert(input_dir, output_dir=None):
    """Batch convert all .html files in a directory to formatted DOCX."""
    input_dir = Path(input_dir)
    if output_dir is None:
        output_dir = input_dir
    else:
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)

    html_files = sorted(input_dir.glob("*.html"))
    if not html_files:
        print("No .html files found in {}".format(input_dir))
        return []

    print("Found {} HTML file(s) to convert...".format(len(html_files)))
    print("-" * 60)

    converted = []
    for html_path in html_files:
        output_path = output_dir / "{}.docx".format(html_path.stem)
        try:
            result = convert_html_to_docx(html_path, output_path)
            if result:
                converted.append(result)
        except Exception as e:
            print("    x Failed: {}: {}".format(html_path.name, e))

    print("-" * 60)
    print("Converted {} of {} file(s)".format(len(converted), len(html_files)))
    return converted


# ======================================================================
# COMMAND LINE INTERFACE
# ======================================================================

def main():
    parser = argparse.ArgumentParser(
        description='Convert notebook HTML files to formatted Word documents',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
    # Auto-convert all .html files in script's folder (double-click / no args)
    python html_to_word.py

    # Single file conversion
    python html_to_word.py notebook.html
    python html_to_word.py notebook.html -o output.docx

    # Batch conversion (all .html files in a directory)
    python html_to_word.py --batch /path/to/html/
    python html_to_word.py --batch /path/to/html/ -o /path/to/output/
        """
    )

    parser.add_argument('input', nargs='?', help='Input .html file (or directory with --batch)')
    parser.add_argument('-o', '--output', help='Output .docx file or directory (for --batch)')
    parser.add_argument('--batch', action='store_true',
                        help='Batch convert all .html files in the input directory')

    args = parser.parse_args()

    print("=" * 60)
    print("    HTML -> Formatted Word Converter")
    print("    (Colors + Spacing + Tables + Images)")
    print("=" * 60)

    if not args.input:
        script_dir = Path(__file__).parent.resolve()
        work_dir = Path.cwd()
        if work_dir != script_dir:
            script_dir = work_dir

        print("\nNo input specified. Auto-converting all .html files in:")
        print("  {}".format(script_dir))

        output_dir = args.output if args.output else str(script_dir / "Word_Outputs")
        batch_convert(input_dir=str(script_dir), output_dir=output_dir)

    elif args.batch:
        batch_convert(input_dir=args.input, output_dir=args.output)

    else:
        input_path = Path(args.input)
        if not input_path.exists():
            print("\n  x File not found: {}".format(input_path))
            return
        output_path = Path(args.output) if args.output else input_path.with_suffix('.docx')
        convert_html_to_docx(input_path, output_path)

    print("\n" + "=" * 60)
    print("  Done!")
    print("=" * 60)


if __name__ == "__main__":
    main()
