"""
Extract a clean, indented Table of Contents (TOC) from a Word (.docx) file.
✓ Keeps page numbers aligned
✓ Removes dotted leaders
✓ Saves as a plain .txt file (ready for Notepad)
"""

import os
import re
import sys
from docx import Document

# === 🪄 EDIT THIS PATH ONLY ===
WORD_FILE = r"C:\Users\manga\OneDrive\####Mind_Palace\####Technical\##AI\3_Machine_learning\Chapter_wise\Chapter_4_5_machine_learning.docx"

# === Output path ===
OUTPUT_FILE = os.path.join(os.path.expanduser("~"), "Desktop", "TOC_Indented.txt")

# === Spaces per indent level ===
INDENT_SPACES = 4

# === Which mode to use: "auto", "headings", or "toc" ===
MODE = "auto"  # "auto" tries both Headings and TOC levels

# === Regex patterns for matching Word styles ===
HEADING_RE = re.compile(r"^Heading\s+(\d+)$", re.IGNORECASE)
TOC_RE = re.compile(r"^TOC\s+(\d+)$", re.IGNORECASE)


def clean_toc_line(text: str, keep_pages: bool = True) -> str:
    """
    Clean a TOC line.
    If keep_pages=True, retains page numbers at the end.
    """
    text = text.strip()

    # If the line has a tab, usually separates title and page number
    if "\t" in text:
        parts = text.split("\t")
        title = parts[0].strip()
        page = parts[-1].strip()
        if keep_pages and page:
            return f"{title} {page}"
        else:
            return title

    # Remove dotted leaders if they exist (e.g. "...... 12")
    text = re.sub(r"\.{3,}\s*\d+\s*$", "", text)
    return text.strip()


def detect_level(style_name: str, mode: str):
    """Return heading or TOC level number."""
    style_name = style_name or ""
    if mode in ("auto", "headings"):
        m = HEADING_RE.match(style_name)
        if m:
            return int(m.group(1))
    if mode in ("auto", "toc"):
        m = TOC_RE.match(style_name)
        if m:
            return int(m.group(1))
    return None


def extract_toc(docx_path: str, out_path: str, indent_spaces: int, mode: str):
    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"❌ File not found:\n{docx_path}")

    try:
        doc = Document(docx_path)
    except Exception as e:
        raise RuntimeError(f"⚠️ Could not open Word file.\n{e}")

    lines = []
    for p in doc.paragraphs:
        style = getattr(p.style, "name", "")
        level = detect_level(style, mode)

        # if all TOC entries are labeled "TOC 2", force them to top level
        if "TOC" in style and level == 2:
            level = 1

        if level is None:
            continue

        text = p.text.strip()
        if not text:
            continue

        if TOC_RE.match(style) or mode == "toc":
            text = clean_toc_line(text, keep_pages=True)
        else:
            text = clean_toc_line(text, keep_pages=False)

        if not text:
            continue

        indent = " " * ((max(level, 1) - 1) * indent_spaces)
        lines.append(f"{indent}{text}")

    if not lines:
        raise RuntimeError("⚠️ No TOC or Heading lines found in this file.")

    # --- Check if the last visible paragraph might hold leftover text ---
    last_p = doc.paragraphs[-1].text.strip()
    if last_p and last_p not in lines[-1]:
        # Avoid duplicates; add if it looks like the missing final entry
        if any(ch.isdigit() for ch in last_p):  # e.g., "Preface 17"
            print(f"ℹ️  Added possible last TOC line: {last_p}")
            lines.append(last_p)

    # --- Check if the last visible paragraph might hold leftover text ---
    last_p = doc.paragraphs[-1].text.strip()
    if last_p and last_p not in lines[-1]:
        # Avoid duplicates; add if it looks like the missing final entry
        if any(ch.isdigit() for ch in last_p):  # e.g., "Preface 17"
            print(f"ℹ️  Added possible last TOC line: {last_p}")
            lines.append(last_p)

    # --- Write to file ---
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    with open(out_path, "w", encoding="utf-8", newline="\n") as f:
        f.write("\n".join(lines) + "\n")



    print(f"✅ TOC exported successfully:\n{out_path}")


if __name__ == "__main__":
    try:
        extract_toc(WORD_FILE, OUTPUT_FILE, INDENT_SPACES, MODE)
    except Exception as e:
        print(f"\n{e}")
        input("\nPress Enter to close...")
        sys.exit(1)

    print("\n🎉 Done! TOC_Indented.txt created on Desktop.")
    input("\nPress Enter to close...")
