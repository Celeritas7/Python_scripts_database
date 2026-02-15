"""
HTML → Formatted Word Converter
================================
Takes custom notebook HTML files and converts them to properly
formatted Word documents via Pandoc + python-docx post-processing.

Fixes:
  - Narrow margins (0.5")
  - Compact paragraph spacing (no blank pages)
  - Images scaled to fit page without overflow
  - Tables formatted with dark headers + zebra stripes
  - Matplotlib artifacts removed
  - Duplicate titles removed
  - Heading styles with custom colors

Usage:
  python html_to_word.py input.html [output.docx]
  python html_to_word.py  (processes all .html in current folder)
"""

import sys
import re
import os
import subprocess
import shutil
from pathlib import Path

# ── Dependencies ──────────────────────────────────────────────────────
if not shutil.which('pandoc'):
    print("ERROR: pandoc not found. Install from https://pandoc.org")
    sys.exit(1)

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Emu, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: pip install python-docx")
    sys.exit(1)


# ══════════════════════════════════════════════════════════════════════
# CONFIGURATION
# ══════════════════════════════════════════════════════════════════════

MARGIN = Inches(0.5)         # Narrow margins
MAX_IMAGE_WIDTH = Inches(6.5)  # Max image width (page - margins)
MAX_IMAGE_HEIGHT = Inches(4.5) # Max image height to prevent page overflow
IMAGE_SCALE = 0.80             # Scale images to 80% after fit

# Heading styles
HEADING_STYLES = {
    'Title':     {'size': 26, 'bold': True,  'color': (0x17, 0x36, 0x5D)},
    'Heading 1': {'size': 16, 'bold': True,  'color': (0x36, 0x5F, 0x91)},
    'Heading 2': {'size': 14, 'bold': True,  'color': (0x4F, 0x81, 0xBD)},
    'Heading 3': {'size': 12, 'bold': True,  'color': (0x4F, 0x81, 0xBD)},
    'Heading 4': {'size': 11, 'bold': True,  'color': (0x4F, 0x81, 0xBD)},
}

# Spacing overrides (in Pt) — these fix the blank page problem
STYLE_SPACING = {
    'Title':           {'before': 6,  'after': 3,  'line': 1.0},
    'Heading 1':       {'before': 12, 'after': 3,  'line': 1.0},
    'Heading 2':       {'before': 10, 'after': 2,  'line': 1.0},
    'Heading 3':       {'before': 8,  'after': 2,  'line': 1.0},
    'Heading 4':       {'before': 8,  'after': 2,  'line': 1.0},
    'Body Text':       {'before': 2,  'after': 2,  'line': 1.0},
    'First Paragraph': {'before': 2,  'after': 2,  'line': 1.0},
    'Compact':         {'before': 1,  'after': 1,  'line': 1.0},
    'Source Code':     {'before': 1,  'after': 1,  'line': 1.0},
    'Block Text':      {'before': 2,  'after': 2,  'line': 1.0},
    'Normal':          {'before': 2,  'after': 2,  'line': 1.0},
}

# Patterns to remove from output
UNWANTED_PATTERNS = [
    r'<Axes:.*?>',
    r'<AxesSubplot:.*?>',
    r'<matplotlib\..*?>',
    r'<Figure.*?>',
    r'<mpl_toolkits\..*?>',
]


# ══════════════════════════════════════════════════════════════════════
# FORMATTING FUNCTIONS
# ══════════════════════════════════════════════════════════════════════

def set_margins(doc):
    """Set narrow margins on all sections."""
    for section in doc.sections:
        section.top_margin = MARGIN
        section.bottom_margin = MARGIN
        section.left_margin = MARGIN
        section.right_margin = MARGIN
        # Set page size to A4 if not set
        if section.page_width is None:
            section.page_width = Inches(8.27)
            section.page_height = Inches(11.69)


def fix_style_spacing(doc):
    """Fix all style spacing to be compact — eliminates blank pages."""
    for style_name, spacing in STYLE_SPACING.items():
        try:
            style = doc.styles[style_name]
            pf = style.paragraph_format
            pf.space_before = Pt(spacing['before'])
            pf.space_after = Pt(spacing['after'])
            # Set line spacing
            if 'line' in spacing:
                pf.line_spacing = spacing['line']
        except KeyError:
            pass  # Style doesn't exist in this doc


def apply_heading_styles(doc):
    """Apply custom heading font styles."""
    for style_name, fmt in HEADING_STYLES.items():
        try:
            style = doc.styles[style_name]
            font = style.font
            font.size = Pt(fmt['size'])
            font.bold = fmt['bold']
            font.color.rgb = RGBColor(*fmt['color'])
        except KeyError:
            pass


def remove_duplicate_title(doc):
    """Remove duplicate title (Pandoc creates both Title + Heading 1)."""
    if len(doc.paragraphs) < 2:
        return 0

    first = doc.paragraphs[0]
    second = doc.paragraphs[1]

    # If first two paragraphs have same text, remove one
    if (first.style and second.style and
        first.text.strip().rstrip('¶') == second.text.strip().rstrip('¶') and
        'Title' in first.style.name):
        p = first._element
        p.getparent().remove(p)
        return 1
    return 0


def remove_unwanted_paragraphs(doc):
    """Remove paragraphs with matplotlib artifacts."""
    removed = 0
    to_remove = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            for pattern in UNWANTED_PATTERNS:
                if re.fullmatch(pattern, text):
                    to_remove.append(para)
                    break

    for para in to_remove:
        try:
            p = para._element
            p.getparent().remove(p)
            removed += 1
        except:
            pass

    return removed


def clean_heading_anchors(doc):
    """Remove trailing ¶ characters from headings."""
    for para in doc.paragraphs:
        if para.style and 'Heading' in para.style.name:
            for run in para.runs:
                if run.text and run.text.endswith('¶'):
                    run.text = run.text.rstrip('¶').rstrip()
            # Also remove hyperlink-based anchors with ¶
            for hl in para._element.findall(qn('w:hyperlink')):
                hl_text = ''.join(r.text or '' for r in hl.findall('.//' + qn('w:t')))
                if hl_text.strip() == '¶':
                    para._element.remove(hl)


def process_images(doc):
    """Scale images to fit within page, prevent overflow that causes blank pages."""
    processed = 0

    for shape in doc.inline_shapes:
        try:
            inline = shape._inline
            extent = inline.find(qn('wp:extent'))
            if extent is None:
                continue

            cx = int(extent.get('cx'))
            cy = int(extent.get('cy'))

            # Calculate max dimensions
            max_w = int(MAX_IMAGE_WIDTH)
            max_h = int(MAX_IMAGE_HEIGHT)

            # Scale to fit within bounds
            new_cx, new_cy = cx, cy

            # First fit to width
            if new_cx > max_w:
                ratio = max_w / new_cx
                new_cx = max_w
                new_cy = int(new_cy * ratio)

            # Then fit to height
            if new_cy > max_h:
                ratio = max_h / new_cy
                new_cy = max_h
                new_cx = int(new_cx * ratio)

            # Apply additional scale factor
            new_cx = int(new_cx * IMAGE_SCALE)
            new_cy = int(new_cy * IMAGE_SCALE)

            # Update extent
            extent.set('cx', str(new_cx))
            extent.set('cy', str(new_cy))

            # Update graphic extent too
            for ext in inline.iter(qn('a:ext')):
                if ext.get('cx') and ext.get('cy'):
                    ext.set('cx', str(new_cx))
                    ext.set('cy', str(new_cy))

            # Center the image paragraph
            parent = inline.getparent()
            while parent is not None:
                if parent.tag == qn('w:p'):
                    pPr = parent.find(qn('w:pPr'))
                    if pPr is None:
                        pPr = OxmlElement('w:pPr')
                        parent.insert(0, pPr)
                    jc = pPr.find(qn('w:jc'))
                    if jc is None:
                        jc = OxmlElement('w:jc')
                        pPr.append(jc)
                    jc.set(qn('w:val'), 'center')
                    break
                parent = parent.getparent()

            processed += 1
        except Exception:
            pass

    return processed


def set_cell_shading(cell, hex_color):
    """Set background color on a table cell."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), hex_color)
    shading.set(qn('w:val'), 'clear')
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn('w:shd'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(shading)


def set_cell_borders(cell, color='000000', size='4'):
    """Set borders on a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), size)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    existing = tcPr.find(qn('w:tcBorders'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(tcBorders)


def format_tables(doc):
    """Format tables: dark header, zebra stripes, compact text, borders."""
    formatted = 0

    for table in doc.tables:
        try:
            # Set table alignment to center
            tblPr = table._tbl.tblPr
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                table._tbl.insert(0, tblPr)

            jc = tblPr.find(qn('w:jc'))
            if jc is None:
                jc = OxmlElement('w:jc')
                tblPr.append(jc)
            jc.set(qn('w:val'), 'center')

            for row_idx, row in enumerate(table.rows):
                for cell in row.cells:
                    # Set borders
                    set_cell_borders(cell, '999999', '4')

                    # Header row: dark background, white text
                    if row_idx == 0:
                        set_cell_shading(cell, '2F2F2F')
                        for para in cell.paragraphs:
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in para.runs:
                                run.font.bold = True
                                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                                run.font.size = Pt(9)
                    else:
                        # Zebra stripes
                        if row_idx % 2 == 0:
                            set_cell_shading(cell, 'F2F2F2')
                        else:
                            set_cell_shading(cell, 'FFFFFF')

                        for para in cell.paragraphs:
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in para.runs:
                                run.font.size = Pt(9)

                    # Compact cell spacing
                    for para in cell.paragraphs:
                        para.paragraph_format.space_before = Pt(1)
                        para.paragraph_format.space_after = Pt(1)

            # Set table-level borders
            tblBorders = OxmlElement('w:tblBorders')
            for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{edge}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '999999')
                tblBorders.append(border)

            existing_borders = tblPr.find(qn('w:tblBorders'))
            if existing_borders is not None:
                tblPr.remove(existing_borders)
            tblPr.append(tblBorders)

            formatted += 1

        except Exception as e:
            print(f"  [WARN] Table format error: {e}")

    return formatted


def reduce_paragraph_spacing(doc):
    """Additional pass: override any per-paragraph spacing that's too large."""
    fixed = 0
    max_before = Pt(14)  # Max 14pt before any paragraph
    max_after = Pt(6)    # Max 6pt after

    for para in doc.paragraphs:
        pf = para.paragraph_format
        changed = False

        if pf.space_before and pf.space_before > max_before:
            pf.space_before = max_before
            changed = True
        if pf.space_after and pf.space_after > max_after:
            pf.space_after = max_after
            changed = True

        if changed:
            fixed += 1

    return fixed


def set_source_code_font(doc):
    """Ensure Source Code paragraphs use monospace font."""
    for para in doc.paragraphs:
        if para.style and para.style.name == 'Source Code':
            for run in para.runs:
                if not run.font.name:
                    run.font.name = 'Consolas'
                if not run.font.size:
                    run.font.size = Pt(9)


# ══════════════════════════════════════════════════════════════════════
# MAIN CONVERSION
# ══════════════════════════════════════════════════════════════════════

def convert_html_to_docx(html_path: Path, output_path: Path) -> bool:
    """Convert HTML → DOCX via Pandoc, then apply all formatting fixes."""
    print(f"\n  Processing: {html_path.name}")

    try:
        # Step 1: Pandoc HTML → DOCX
        print("    [1/3] Converting HTML → DOCX via Pandoc...")
        result = subprocess.run([
            'pandoc',
            '--from=html',
            '--to=docx',
            '--standalone',
            f'--resource-path={str(html_path.parent)}',
            '--wrap=none',
            '-o', str(output_path),
            str(html_path),
        ], capture_output=True, text=True)
        if result.returncode != 0:
            print(f"    [ERROR] Pandoc failed: {result.stderr}")
            return False

        # Step 2: Open and fix
        print("    [2/3] Fixing spacing & layout...")
        doc = Document(str(output_path))

        set_margins(doc)
        fix_style_spacing(doc)
        apply_heading_styles(doc)

        dup = remove_duplicate_title(doc)
        artifacts = remove_unwanted_paragraphs(doc)
        clean_heading_anchors(doc)
        spacing_fixed = reduce_paragraph_spacing(doc)

        # Step 3: Format content
        print("    [3/3] Formatting tables & images...")
        tables = format_tables(doc)
        images = process_images(doc)
        set_source_code_font(doc)

        # Save
        doc.save(str(output_path))

        print(f"      Tables: {tables} | Images: {images} | "
              f"Artifacts removed: {artifacts} | Spacing fixes: {spacing_fixed}")
        print(f"    [DONE] → {output_path.name}")
        return True

    except Exception as e:
        print(f"    [ERROR] {e}")
        import traceback
        traceback.print_exc()
        return False


def main():
    print("=" * 60)
    print("    HTML → Formatted Word Converter")
    print("=" * 60)

    args = sys.argv[1:]

    if args:
        # Specific file(s)
        html_files = [Path(a) for a in args if a.endswith('.html')]
        if not html_files:
            # If a docx output was specified, separate it
            html_files = [Path(args[0])]
    else:
        # All HTML in current directory
        html_files = list(Path('.').glob('*.html'))

    if not html_files:
        print("\n  No .html files found.")
        print("  Usage: python html_to_word.py input.html [output.docx]")
        return

    print(f"\n  Found {len(html_files)} file(s)")

    output_dir = Path('Word_Outputs')
    output_dir.mkdir(exist_ok=True)

    success = 0
    for html_path in html_files:
        out_path = output_dir / f"{html_path.stem}.docx"
        if convert_html_to_docx(html_path, out_path):
            success += 1

    print(f"\n{'=' * 60}")
    print(f"  Complete! {success}/{len(html_files)} converted → {output_dir}/")
    print(f"{'=' * 60}")


if __name__ == "__main__":
    main()
