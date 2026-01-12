import argparse
import os
import re
import sys
from docx import Document

HEADING_RE = re.compile(r'^Heading\s+(\d+)$', re.IGNORECASE)
TOC_RE      = re.compile(r'^TOC\s+(\d+)$', re.IGNORECASE)

def clean_toc_line(text: str) -> str:
    """
    Removes page numbers and dot leaders often found in TOC lines.
    Examples:
      'Chapter 1.............  12' -> 'Chapter 1'
      'Intro\t3'                      -> 'Intro'
    """
    # If there's a hard tab (common for TOC leaders), keep only text before it
    if '\t' in text:
        text = text.split('\t', 1)[0]
    # Remove dot leader + page number at the end, e.g. ".......  12"
    text = re.sub(r'\.{3,}\s*\d+\s*$', '', text)
    # Remove trailing page numbers if they appear without dots
    text = re.sub(r'\s+\d+\s*$', '', text)
    return text.strip()

def detect_level(style_name: str, mode: str):
    """
    Returns an integer level or None.
    mode:
      - 'auto': try headings first, then TOC
      - 'headings': only Heading N
      - 'toc': only TOC N
    """
    style_name = style_name or ""
    m = None

    if mode in ("auto", "headings"):
        m = HEADING_RE.match(style_name)
        if m:
            return int(m.group(1))

    if mode in ("auto", "toc"):
        m = TOC_RE.match(style_name)
        if m:
            return int(m.group(1))

    return None

def extract(docx_path: str, out_path: str, indent_spaces: int, mode: str):
    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"Input file not found: {docx_path}")

    # Ensure OneDrive file is locally available (files-on-demand can be cloud-only)
    try:
        doc = Document(docx_path)
    except Exception as e:
        raise RuntimeError(f"Failed to open DOCX. If it's on OneDrive, make sure it's downloaded locally.\n{e}")

    lines = []
    for p in doc.paragraphs:
        style = getattr(p.style, "name", "") or ""
        level = detect_level(style, mode)
        if level is None:
            continue

        text = p.text.strip()
        if not text:
            continue

        # If this is a TOC line, clean leaders/page numbers
        if TOC_RE.match(style) or mode == "toc":
            text = clean_toc_line(text)

        # Skip empty after cleaning
        if not text:
            continue

        indent = " " * max(0, (level - 1) * indent_spaces)
        lines.append(f"{indent}{text}")

    if not lines and mode != "toc":
        # As a fallback: try to parse visible TOC-only lines if headings gave nothing
        for p in doc.paragraphs:
            style = getattr(p.style, "name", "") or ""
            m = TOC_RE.match(style)
            if not m:
                continue
            level = int(m.group(1))
            text = clean_toc_line(p.text)
            if text:
                indent = " " * max(0, (level - 1) * indent_spaces)
                lines.append(f"{indent}{text}")

    # Final check
    if not lines:
        raise RuntimeError(
            "No headings/TOC lines found.\n"
            "- If you used Word’s auto TOC: set mode to --mode toc\n"
            "- If you have headings: ensure they use 'Heading 1/2/3' styles\n"
            "- If the file is synced (OneDrive), ensure it's downloaded locally"
        )

    # Write output
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    with open(out_path, "w", encoding="utf-8", newline="\n") as f:
        f.write("\n".join(lines) + "\n")

def main():
    parser = argparse.ArgumentParser(
        description="Extract an indented plain-text TOC from a .docx (headings or TOC styles)."
    )
    parser.add_argument("docx", help="C:\Users\manga\OneDrive\####Mind_Palace\####Technical\##AI\3_Machine_learning\Index.docx")
    parser.add_argument("-o", "--out", default=os.path.join(os.path.expanduser("~"), "Desktop", "TOC_Indented.txt"),
                        help="Output .txt path (default: Desktop\\TOC_Indented.txt)")
    parser.add_argument("--indent", type=int, default=4, help="Spaces per level (default: 4)")
    parser.add_argument("--mode", choices=["auto", "headings", "toc"], default="auto",
                        help="Parse mode: 'headings' (Heading 1/2/3), 'toc' (TOC 1/2/3), or 'auto' (default)")
    args = parser.parse_args()

    try:
        extract(args.docx, args.out, args.indent, args.mode)
        print(f"✓ Saved: {args.out}")
    except Exception as e:
        print(f"ERROR: {e}", file=sys.stderr)
        sys.exit(1)

if __name__ == "__main__":
    main()
