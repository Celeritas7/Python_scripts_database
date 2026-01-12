from docx import Document
from docx.enum.text import WD_COLOR_INDEX
import re

# Load the document
doc = Document(r"C:\Users\manga\OneDrive\####Mind_Palace\####Technical\##AI\Calculus_testing_1.docx")

# Regex to capture expressions like H(t)=...
formula_pattern = r'[A-Za-z]+\([^)]*\)=−?\d+[^\s]*'

seen = set()

for para in doc.paragraphs:
    text = para.text
    matches = re.finditer(formula_pattern, text)

    new_runs = []
    last_index = 0

    for match in matches:
        expr = match.group()
        norm_expr = expr.replace("−", "-").replace(" ", "")  # Normalize minus and spacing
        start, end = match.span()

        # Add unhighlighted text before match
        if start > last_index:
            new_runs.append((text[last_index:start], None))

        # Highlight if duplicate
        if norm_expr in seen:
            new_runs.append((text[start:end], WD_COLOR_INDEX.YELLOW))
        else:
            new_runs.append((text[start:end], None))
            seen.add(norm_expr)

        last_index = end

    # Add remaining text
    if last_index < len(text):
        new_runs.append((text[last_index:], None))

    # Rebuild paragraph
    if new_runs:
        para.clear()
        for run_text, highlight in new_runs:
            run = para.add_run(run_text)
            if highlight:
                run.font.highlight_color = highlight

# Save the updated document
doc.save(r"C:\Users\manga\OneDrive\####Mind_Palace\####Technical\##AI\Calculus_testing_2.docx")